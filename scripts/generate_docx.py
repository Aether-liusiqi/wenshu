#!/usr/bin/env python3
"""Word (.docx) exporter for Chinese official documents (公文 Word 导出器).

Reads a Markdown document and exports it to a properly formatted .docx file
with A4 page, correct fonts, heading styles, and layout.

Requires: python-docx (pip install python-docx)
"""

import re
import sys
from pathlib import Path
from dataclasses import dataclass
from typing import Optional

try:
    from docx import Document
    from docx.shared import Pt, Cm, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.section import WD_ORIENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("错误: 需要 python-docx 库。请运行: pip install python-docx")
    sys.exit(1)


# ── Font profiles ──────────────────────────────────────────────
# Each profile maps Markdown elements to (font_name, font_size_pt, bold, italic)

@dataclass
class FontSpec:
    name: str
    name_east: str  # East Asian font
    size: int       # points
    bold: bool = False
    italic: bool = False
    color: Optional[str] = None  # hex RGB


FONT_PROFILES: dict[str, dict] = {
    "standard": {
        "title": FontSpec("Times New Roman", "小标宋体", 22, bold=False),
        "h1": FontSpec("Times New Roman", "黑体", 16, bold=True),
        "h2": FontSpec("Times New Roman", "楷体", 16, bold=True),
        "body": FontSpec("Times New Roman", "仿宋体", 16, bold=False),
        "footer": FontSpec("Times New Roman", "仿宋体", 14, bold=False),
    },
    "compact": {  # for reports, briefings
        "title": FontSpec("Times New Roman", "小标宋体", 18, bold=False),
        "h1": FontSpec("Times New Roman", "黑体", 14, bold=True),
        "h2": FontSpec("Times New Roman", "楷体", 14, bold=True),
        "body": FontSpec("Times New Roman", "仿宋体", 14, bold=False),
        "footer": FontSpec("Times New Roman", "仿宋体", 12, bold=False),
    },
}

# ── Layout profiles ─────────────────────────────────────────────

@dataclass
class LayoutSpec:
    paper_size: tuple[float, float]  # (width_cm, height_cm)
    margin_top: float     # cm
    margin_bottom: float  # cm
    margin_left: float    # cm
    margin_right: float   # cm
    body_indent: float    # cm (first line indent)
    line_spacing: float   # multiplier
    body_spacing_after: float  # pt
    show_page_number: bool = True


LAYOUT_PROFILES: dict[str, LayoutSpec] = {
    "standard": LayoutSpec(
        paper_size=(21.0, 29.7),
        margin_top=3.7,
        margin_bottom=3.5,
        margin_left=2.8,
        margin_right=2.6,
        body_indent=0.74,  # 2 chars at 3号 font
        line_spacing=1.5,
        body_spacing_after=0,
        show_page_number=True,
    ),
    "compact": LayoutSpec(
        paper_size=(21.0, 29.7),
        margin_top=2.5,
        margin_bottom=2.5,
        margin_left=2.5,
        margin_right=2.0,
        body_indent=0.7,
        line_spacing=1.3,
        body_spacing_after=2,
        show_page_number=True,
    ),
}

# ── Doc type → profile mapping ──────────────────────────────────

DOC_TYPE_PROFILES: dict[str, tuple[str, str]] = {
    # (font_profile, layout_profile)
    "通知": ("standard", "standard"),
    "报告": ("standard", "standard"),
    "请示": ("standard", "standard"),
    "函": ("standard", "standard"),
    "纪要": ("standard", "standard"),
    "决定": ("standard", "standard"),
    "通报": ("standard", "standard"),
    "公告": ("standard", "standard"),
    "通告": ("standard", "standard"),
    "意见": ("standard", "standard"),
    "批复": ("standard", "standard"),
    "议案": ("standard", "standard"),
    "决议": ("standard", "standard"),
    "公报": ("standard", "standard"),
    "命令": ("standard", "standard"),
    # 事务文书
    "工作总结": ("compact", "compact"),
    "工作方案": ("compact", "compact"),
    "讲话稿": ("compact", "compact"),
    "汇报材料": ("compact", "compact"),
    "简报": ("compact", "compact"),
    "情况专报": ("compact", "compact"),
    "回复函": ("standard", "standard"),
}


def load_markdown(filepath: str) -> str:
    path = Path(filepath)
    if not path.exists():
        raise FileNotFoundError(f"文件不存在: {filepath}")
    return path.read_text(encoding="utf-8")


def parse_markdown_structure(text: str) -> list[dict]:
    """Parse Markdown into a list of structural elements."""
    elements = []
    lines = text.split("\n")
    i = 0

    while i < len(lines):
        line = lines[i].strip()

        # Skip frontmatter
        if i == 0 and line == "---":
            i += 1
            while i < len(lines) and lines[i].strip() != "---":
                i += 1
            i += 1
            continue

        # Heading
        m = re.match(r"^(#{1,6})\s+(.+)$", line)
        if m:
            level = len(m.group(1))
            heading_text = m.group(2).strip()
            elements.append({"type": "heading", "level": level, "text": heading_text})
            i += 1
            continue

        # Horizontal rule
        if re.match(r"^[-–—]{3,}$", line) or re.match(r"^\*{3,}$", line):
            elements.append({"type": "hr"})
            i += 1
            continue

        # Empty line
        if not line:
            elements.append({"type": "blank"})
            i += 1
            continue

        # Paragraph — collect consecutive non-blank, non-special lines
        para_lines = []
        while i < len(lines):
            ln = lines[i].strip()
            if not ln:
                break
            if re.match(r"^(#{1,6})\s+", ln):
                break
            if re.match(r"^[-–—]{3,}$", ln) or re.match(r"^\*{3,}$", ln):
                break
            # Skip HTML comments and style blocks
            if ln.startswith("<!--") or ln.startswith("<style"):
                i += 1
                continue
            para_lines.append(lines[i].rstrip())
            i += 1

        if para_lines:
            text_body = "\n".join(para_lines).strip()
            # Detect blockquote
            is_blockquote = all(
                ln.strip().startswith(">") for ln in para_lines if ln.strip()
            )
            # Detect table
            is_table = any("|" in ln for ln in para_lines)

            elements.append({
                "type": "blockquote" if is_blockquote else "table" if is_table else "paragraph",
                "text": text_body,
                "raw_lines": para_lines,
            })
        else:
            i += 1

    return elements


def apply_font(run, spec: FontSpec):
    """Apply font specification to a docx run."""
    run.font.name = spec.name
    run.font.size = Pt(spec.size)
    run.bold = spec.bold
    run.italic = spec.italic
    if spec.color:
        run.font.color.rgb = RGBColor(
            int(spec.color[0:2], 16),
            int(spec.color[2:4], 16),
            int(spec.color[4:6], 16),
        )
    # Set East Asian font
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), spec.name_east)


def add_page_number(doc):
    """Add page numbers to document footer."""
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add page number field
        run = p.add_run()
        fld_char_begin = OxmlElement('w:fldChar')
        fld_char_begin.set(qn('w:fldCharType'), 'begin')
        run._element.append(fld_char_begin)

        instr_text = OxmlElement('w:instrText')
        instr_text.set(qn('xml:space'), 'preserve')
        instr_text.text = ' PAGE '
        run._element.append(instr_text)

        fld_char_end = OxmlElement('w:fldChar')
        fld_char_end.set(qn('w:fldCharType'), 'end')
        run._element.append(fld_char_end)

        run.font.size = Pt(9)


def export_docx(
    filepath: str,
    output_path: str,
    doc_type: str = "通知",
    hide_page_number: bool = False,
) -> str:
    """Export a Markdown document to Word .docx format.

    Args:
        filepath: Path to source Markdown file
        output_path: Path to output .docx file
        doc_type: Document type (determines font and layout profiles)
        hide_page_number: If True, suppress page numbers

    Returns:
        Path to the generated .docx file
    """
    text = load_markdown(filepath)
    elements = parse_markdown_structure(text)

    # Select profiles
    font_key, layout_key = DOC_TYPE_PROFILES.get(
        doc_type, ("standard", "standard")
    )
    # Try fuzzy match
    if doc_type not in DOC_TYPE_PROFILES:
        for key in DOC_TYPE_PROFILES:
            if key in doc_type or doc_type in key:
                font_key, layout_key = DOC_TYPE_PROFILES[key]
                break

    fonts = FONT_PROFILES.get(font_key, FONT_PROFILES["standard"])
    layout = LAYOUT_PROFILES.get(layout_key, LAYOUT_PROFILES["standard"])

    # Create document
    doc = Document()

    # Page setup
    for section in doc.sections:
        section.page_width = Cm(layout.paper_size[0])
        section.page_height = Cm(layout.paper_size[1])
        section.top_margin = Cm(layout.margin_top)
        section.bottom_margin = Cm(layout.margin_bottom)
        section.left_margin = Cm(layout.margin_left)
        section.right_margin = Cm(layout.margin_right)

    first_heading_seen = False

    for elem in elements:
        if elem["type"] == "blank":
            continue

        elif elem["type"] == "heading":
            heading_text = elem["text"]
            # Skip comment-style headings (like "格式提示")
            if re.match(r"^(格式提示|文种说明|范文|适用场景|格式要点)", heading_text):
                continue
            # Skip markdown artifact headings
            if heading_text in ("---",):
                continue

            level = min(elem["level"], 3)  # Cap at h3

            if not first_heading_seen:
                # This is the document title
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(heading_text)
                apply_font(run, fonts["title"])
                first_heading_seen = True
            else:
                p = doc.add_paragraph()
                if level == 1:
                    run = p.add_run(heading_text)
                    apply_font(run, fonts["h1"])
                elif level == 2:
                    run = p.add_run(heading_text)
                    apply_font(run, fonts["h2"])
                else:
                    run = p.add_run(heading_text)
                    apply_font(run, fonts["body"])
                    run.bold = True

        elif elem["type"] == "hr":
            # Add a thin separator
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)

        elif elem["type"] == "paragraph":
            raw_text = elem["text"]

            # Strip comment-style annotations
            raw_text = re.sub(r"^>.*?提示.*$", "", raw_text, flags=re.MULTILINE)
            raw_text = re.sub(r"^>.*?文种说明.*$", "", raw_text, flags=re.MULTILINE)

            # Skip format annotation blocks
            if re.match(r"^(\*\*文种\*\*|\*\*格式要点\*\*)", raw_text):
                continue
            if raw_text.startswith("---") and len(raw_text) < 20:
                continue

            # Remove markdown bold/italic markers for cleaner export
            clean_text = re.sub(r"\*\*(.+?)\*\*", r"\1", raw_text)
            clean_text = re.sub(r"\*(.+?)\*", r"\1", clean_text)
            clean_text = re.sub(r"~~(.+?)~~", r"\1", clean_text)
            clean_text = re.sub(r"`(.+?)`", r"\1", clean_text)

            # Skip lines that are pure formatting guides
            clean_lines = []
            for ln in clean_text.split("\n"):
                ln = ln.strip()
                if not ln:
                    clean_lines.append("")
                    continue
                # Skip reference-style links
                if re.match(r"^\[.+\]:\s*http", ln):
                    continue
                clean_lines.append(ln)

            clean_text = "\n".join(clean_lines).strip()
            if not clean_text:
                continue

            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Cm(layout.body_indent)
            p.paragraph_format.line_spacing = layout.line_spacing
            p.paragraph_format.space_after = Pt(layout.body_spacing_after)

            run = p.add_run(clean_text)
            apply_font(run, fonts["body"])

        elif elem["type"] == "blockquote":
            raw_text = elem["text"]
            # Strip the '>' markers
            clean_text = "\n".join(
                re.sub(r"^>\s?", "", ln.strip())
                for ln in raw_text.split("\n")
                if ln.strip().startswith(">")
            ).strip()

            if not clean_text:
                continue

            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1.0)
            p.paragraph_format.line_spacing = layout.line_spacing

            run = p.add_run(clean_text)
            apply_font(run, fonts["body"])
            run.italic = True

        elif elem["type"] == "table":
            # Simple table handling
            raw_lines = elem.get("raw_lines", [])
            # Filter to lines with pipe separators
            table_lines = [ln for ln in raw_lines if "|" in ln and ln.strip().startswith("|")]
            if len(table_lines) < 2:
                # Not a proper table, render as paragraph
                p = doc.add_paragraph()
                p.paragraph_format.first_line_indent = Cm(layout.body_indent)
                run = p.add_run(elem["text"])
                apply_font(run, fonts["body"])
                continue

            # Skip separator line (---|---)
            data_lines = [
                ln for ln in table_lines
                if not re.match(r"^\|[\s\-:|]+\|$", ln.strip())
            ]

            if not data_lines:
                continue

            # Parse rows
            rows = []
            for ln in data_lines:
                cells = [c.strip() for c in ln.strip().strip("|").split("|")]
                rows.append(cells)

            if not rows:
                continue

            ncols = max(len(r) for r in rows)
            table = doc.add_table(rows=len(rows), cols=ncols)
            table.style = 'Table Grid'

            for i, row_data in enumerate(rows):
                for j, cell_text in enumerate(row_data):
                    if j < ncols:
                        cell = table.rows[i].cells[j]
                        cell.text = cell_text
                        for paragraph in cell.paragraphs:
                            paragraph.paragraph_format.space_before = Pt(2)
                            paragraph.paragraph_format.space_after = Pt(2)
                            for run in paragraph.runs:
                                run.font.size = Pt(10)
                                run.font.name = "仿宋体"

    # Add page numbers
    if layout.show_page_number and not hide_page_number:
        add_page_number(doc)

    # Ensure output directory exists
    output = Path(output_path)
    output.parent.mkdir(parents=True, exist_ok=True)

    doc.save(str(output))
    return str(output.resolve())


def main():
    # Force UTF-8 output on Windows
    if hasattr(sys.stdout, 'reconfigure'):
        sys.stdout.reconfigure(encoding='utf-8')

    if len(sys.argv) < 3:
        print("用法: python generate_docx.py <输入.md> -o <输出.docx> [--doc-type 文种] [--hide-page-number]")
        print("示例: python generate_docx.py draft.md -o output.docx --doc-type 通知")
        print()
        print("支持的文种:")
        for k in DOC_TYPE_PROFILES:
            print(f"  - {k}")
        sys.exit(1)

    filepath = sys.argv[1]
    output_path = None
    doc_type = "通知"
    hide_page_number = False

    i = 2
    while i < len(sys.argv):
        if sys.argv[i] == "-o" and i + 1 < len(sys.argv):
            output_path = sys.argv[i + 1]
            i += 2
        elif sys.argv[i] == "--doc-type" and i + 1 < len(sys.argv):
            doc_type = sys.argv[i + 1]
            i += 2
        elif sys.argv[i] == "--hide-page-number":
            hide_page_number = True
            i += 1
        else:
            i += 1

    if not output_path:
        base = Path(filepath).stem
        output_path = f"{base}.docx"

    try:
        result = export_docx(filepath, output_path, doc_type, hide_page_number)
        print(f"✓ Word 文件已导出: {result}")
        print(f"  文种方案: {doc_type}")
    except FileNotFoundError as e:
        print(f"错误: {e}")
        sys.exit(1)
    except Exception as e:
        print(f"导出失败: {e}")
        sys.exit(1)


if __name__ == "__main__":
    main()
